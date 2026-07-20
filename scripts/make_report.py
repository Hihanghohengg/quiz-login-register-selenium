from __future__ import annotations

from pathlib import Path
from datetime import date
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Laporan_Pengujian_Selenium_Login_Register.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")



def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def compact_code(code: str) -> str:
    return "\n".join(line for line in code.splitlines() if line.strip())


def add_code(doc: Document, code: str, title: str | None = None) -> None:
    if title:
        p = doc.add_paragraph()
        p.style = "Code Caption"
        p.add_run(title)
    for line in code.rstrip().splitlines():
        p = doc.add_paragraph(style="Code Block")
        p.add_run(line if line else " ")


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    prevent_row_split(table.rows[0])
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        shade_cell(hdr[i], "1F4E78")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(font_size)
        set_cell_margins(hdr[i])
    for ridx, row in enumerate(rows):
        new_row = table.add_row()
        prevent_row_split(new_row)
        cells = new_row.cells
        if ridx % 2:
            for c in cells:
                shade_cell(c, "EAF2F8")
        for i, text in enumerate(row):
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(font_size)
            set_cell_margins(cells[i])
    return table


def add_status_box(doc: Document, title: str, text: str, fill: str = "D9EAD3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    prevent_row_split(t.rows[0])
    c = t.cell(0, 0)
    shade_cell(c, fill)
    p = c.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    p.add_run(text)
    set_cell_margins(c, top=140, start=160, bottom=140, end=160)


def configure_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 25, RGBColor(31, 78, 120)),
        ("Heading 1", 17, RGBColor(31, 78, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11, RGBColor(55, 95, 145)),
    ]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color
        styles[name].font.bold = True

    if "Code Block" not in styles:
        st = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles["Code Block"]
    st.font.name = "Liberation Mono"
    st.font.size = Pt(6.5)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.left_indent = Inches(0.12)
    st.paragraph_format.right_indent = Inches(0.05)
    st.paragraph_format.keep_together = True
    ppr = st.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    ppr.append(shd)

    if "Code Caption" not in styles:
        cap = styles.add_style("Code Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Code Caption"]
    cap.font.name = "Aptos"
    cap.font.size = Pt(9)
    cap.font.bold = True
    cap.font.color.rgb = RGBColor(31, 78, 120)
    cap.paragraph_format.space_before = Pt(7)
    cap.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        section.header.paragraphs[0].text = "Laporan Pengujian Selenium - Login & Register"
        section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        section.header.paragraphs[0].runs[0].font.size = Pt(8)
        section.header.paragraphs[0].runs[0].font.color.rgb = RGBColor(100, 100, 100)
        add_page_number(section.footer.paragraphs[0])
    return doc


doc = configure_doc()

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(75)
r = p.add_run("LAPORAN PENGUJIAN PERANGKAT LUNAK")
r.bold = True
r.font.size = Pt(25)
r.font.color.rgb = RGBColor(31, 78, 120)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("Otomasi Selenium untuk Modul Login dan Register")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(46, 116, 181)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(25)
p.add_run("Objek uji: hermanka/quiz-pengupil\n").bold = True
p.add_run("Teknologi: PHP, MySQL, Python, pytest, Selenium WebDriver, GitHub Actions")

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = "Table Grid"
for i, (a,b) in enumerate([
    ("Nama", "[ISI NAMA]"),
    ("NIM", "[ISI NIM]"),
    ("Kelas", "[ISI KELAS]"),
    ("Repository hasil", "https://github.com/USERNAME_ANDA/quiz-login-register-selenium"),
    ("Tanggal", "20 Juli 2026"),
]):
    meta.cell(i,0).text=a
    meta.cell(i,1).text=b
    shade_cell(meta.cell(i,0), "D9EAF7")
    meta.cell(i,0).paragraphs[0].runs[0].bold=True
    for c in meta.rows[i].cells:
        set_cell_margins(c, top=120, start=140, bottom=120, end=140)

doc.add_paragraph()
add_status_box(doc, "Catatan sebelum dikumpulkan", "Ganti USERNAME_ANDA, nama, NIM, dan kelas. Setelah paket ZIP diunggah ke GitHub, pastikan workflow di tab Actions berhasil, lalu masukkan URL repository aktual ke PDF.", "FFF2CC")
doc.add_page_break()

# TOC manual
h=doc.add_heading("Daftar Isi", level=1)
sections = [
"1. Ringkasan Eksekutif", "2. Analisis Aplikasi Asal", "3. Persiapan Lingkungan", "4. Strategi dan Arsitektur Pengujian", "5. Daftar Testcase", "6. Implementasi Selenium", "7. Menjalankan Pengujian Lokal", "8. CI/CD GitHub Actions", "9. Mengunggah Repository", "10. Temuan Defect", "11. Validasi dan Checklist Pengumpulan", "12. Kesimpulan", "Lampiran A-C: Kode Utama"
]
for item in sections:
    doc.add_paragraph(item)
doc.add_page_break()

# 1

doc.add_heading("1. Ringkasan Eksekutif", level=1)
doc.add_paragraph(
    "Tugas ini menguji dua modul aplikasi web, yaitu login dan register. Solusi mengotomasi interaksi pengguna melalui browser nyata menggunakan Selenium WebDriver dan pytest. Database MySQL di-reset sebelum setiap testcase agar hasil dapat diulang. Pipeline GitHub Actions menjalankan seluruh suite secara otomatis pada setiap push, pull request, atau pemicu manual."
)
add_table(doc, ["Komponen", "Implementasi"], [
    ["Objek uji", "login.php dan register.php dari repository quiz-pengupil"],
    ["Bahasa otomasi", "Python 3.12"],
    ["Framework", "pytest 8.4.2"],
    ["Browser automation", "Selenium 4.46.0 + ChromeDriver"],
    ["Database", "MySQL 8.4 service container pada GitHub Actions"],
    ["Pola desain", "Page Object Model"],
    ["Stub", "tests/stubs/index.php sebagai pengganti dashboard yang tidak tersedia"],
    ["Jumlah testcase", "17 testcase: 13 normal dan 4 expected failure"],
    ["Bukti hasil", "JUnit XML, log PHP, dan screenshot otomatis ketika testcase gagal"],
], font_size=9)

doc.add_paragraph("Empat testcase diberi penanda xfail(strict=True). Status tersebut bukan berarti script Selenium rusak; penanda itu merekam perilaku salah yang memang ada pada aplikasi asal sehingga pipeline tetap memberikan laporan yang jujur dan stabil.")

# 2

doc.add_heading("2. Analisis Aplikasi Asal", level=1)
doc.add_heading("2.1 Struktur repository", level=2)
add_bullets(doc, [
    "login.php: form username/password dan autentikasi menggunakan tabel users.",
    "register.php: form nama, email, username, password, dan re-password.",
    "koneksi.php: koneksi mysqli ke database quiz_pengupil.",
    "db/quiz_pengupil.sql: skema tabel users dan data contoh.",
    "style.css: tampilan sederhana berbasis Bootstrap.",
    "Repo asal tidak menyediakan index.php walaupun login dan register mengarah ke file tersebut setelah berhasil.",
])

doc.add_heading("2.2 Alur login", level=2)
add_numbered(doc, [
    "Pengguna membuka login.php.",
    "Pengguna memasukkan username dan password.",
    "Server memastikan kedua field tidak kosong.",
    "Server mencari record berdasarkan username.",
    "Server memverifikasi password menggunakan password_verify.",
    "Jika valid, session username dibuat dan pengguna diarahkan ke index.php.",
])

doc.add_heading("2.3 Alur register", level=2)
add_numbered(doc, [
    "Pengguna membuka register.php.",
    "Pengguna mengisi nama, email, username, password, dan konfirmasi password.",
    "Server menolak data kosong dan password yang tidak sama.",
    "Server seharusnya memeriksa username duplikat.",
    "Password di-hash lalu data pengguna disimpan.",
    "Jika sukses, session username dibuat dan pengguna diarahkan ke index.php.",
])

doc.add_heading("2.4 Defect awal yang memengaruhi desain test", level=2)
add_table(doc, ["ID", "Defect", "Dampak"], [
    ["DEF-01", "Login password salah tidak mengisi variabel error.", "Pengguna tetap di halaman login tanpa penjelasan."],
    ["DEF-02", "Username login tidak ditemukan menampilkan 'Register User Gagal'.", "Pesan salah konteks dan membingungkan."],
    ["DEF-03", "Register memanggil cek_nama($name), bukan cek_nama($username).", "Username duplikat dapat lolos."],
    ["DEF-04", "INSERT register menggunakan $nama yang tidak didefinisikan.", "Kolom name dapat kosong atau query gagal, tergantung konfigurasi PHP/MySQL."],
    ["DEF-05", "Tidak ada UNIQUE constraint pada username/email.", "Integritas data duplikat tidak dijaga oleh database."],
    ["DEF-06", "index.php tidak tersedia.", "Positive flow tidak dapat diverifikasi tanpa stub."],
], font_size=8.5)

# 3

doc.add_heading("3. Persiapan Lingkungan", level=1)
doc.add_heading("3.1 Yang harus disiapkan", level=2)
add_bullets(doc, [
    "Akun GitHub dan Git.",
    "PHP 8.x dengan ekstensi mysqli.",
    "MySQL atau MariaDB. Pada Windows paling mudah memakai XAMPP.",
    "Python 3.11 atau lebih baru.",
    "Google Chrome atau Chromium.",
    "ChromeDriver yang kompatibel. Pada runner GitHub Ubuntu, browser dan driver sudah tersedia; script juga mendukung Selenium Manager jika driver tidak ditemukan.",
    "Editor seperti Visual Studio Code.",
])

doc.add_heading("3.2 Persiapan Windows dengan XAMPP", level=2)
add_numbered(doc, [
    "Instal XAMPP lalu aktifkan Apache dan MySQL dari XAMPP Control Panel.",
    "Salin folder proyek ke lokasi kerja biasa; PHP built-in server akan dipakai agar path konsisten.",
    "Buka phpMyAdmin atau command prompt MySQL dan impor db/quiz_pengupil.sql.",
    "Instal Python, centang opsi Add Python to PATH.",
    "Pastikan perintah php, python, pip, git, chrome, dan chromedriver dapat ditemukan.",
])

add_code(doc, """php --version
python --version
git --version
chromedriver --version""", "Perintah pengecekan versi")

doc.add_heading("3.3 Virtual environment", level=2)
add_code(doc, """python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt""", "Windows PowerShell")
add_code(doc, """python -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt""", "Linux/macOS")

# 4

doc.add_heading("4. Strategi dan Arsitektur Pengujian", level=1)
doc.add_heading("4.1 Ruang lingkup", level=2)
add_bullets(doc, [
    "Functional testing terhadap input, validasi, navigasi, session, redirect, dan penyimpanan database.",
    "Negative testing untuk data kosong, password salah, password tidak sama, email salah, username duplikat, dan SQL injection sederhana.",
    "UI smoke testing untuk keberadaan field, tombol, link, dan masking password.",
    "Tidak menguji dashboard karena tidak tersedia dan berada di luar ruang lingkup dua modul.",
])

doc.add_heading("4.2 Arsitektur", level=2)
add_table(doc, ["Lapisan", "Tanggung jawab"], [
    ["Test case", "Menyatakan skenario, data uji, aksi, dan assertion."],
    ["Page Object", "Menyimpan locator dan operasi halaman agar test tidak duplikatif."],
    ["WebDriver", "Mengendalikan Chrome secara headless maupun non-headless."],
    ["Database fixture", "TRUNCATE dan seed user sebelum setiap test."],
    ["Index stub", "Menerima redirect sukses tanpa memerlukan dashboard asli."],
    ["GitHub Actions", "Menyediakan MySQL, PHP server, browser, driver, eksekusi test, dan artifact."],
], font_size=9)

doc.add_heading("4.3 Mengapa Stub dan Driver dipakai", level=2)
doc.add_paragraph("ChromeDriver adalah driver konkret yang menerjemahkan perintah Selenium menjadi aksi browser. Stub index.php dipakai karena tujuan redirect tidak ada pada repo asal. Stub hanya menampilkan marker TEST_STUB_INDEX dan username session; ia tidak meniru fitur bisnis lain sehingga pengujian tetap terfokus.")

doc.add_heading("4.4 Isolasi data", level=2)
doc.add_paragraph("Fixture reset_database berjalan otomatis sebelum setiap testcase. Tabel users dikosongkan lalu user deterministik testuser/Password123! ditambahkan kembali. Dengan demikian, urutan eksekusi testcase tidak memengaruhi hasil.")

# 5 Test cases

doc.add_heading("5. Daftar Testcase", level=1)
doc.add_paragraph("Kolom Status Otomasi menunjukkan ekspektasi terhadap aplikasi asal. Pass berarti testcase seharusnya lulus. XFail berarti assertion yang benar sengaja gagal karena defect yang telah diketahui.")

login_rows = [
["LGN-001", "Elemen halaman tampil", "Buka /login.php", "Judul Sign-In, username, password, tombol terlihat", "Pass"],
["LGN-002", "Navigasi ke register", "Klik link Register", "URL register.php dan judul Sign-Up", "Pass"],
["LGN-003", "Field kosong", "Submit username/password kosong", "Pesan Data tidak boleh kosong", "Pass"],
["LGN-004", "Kredensial valid", "testuser / Password123!", "Redirect index.php, marker stub dan username tampil", "Pass"],
["LGN-005", "Username tidak ada", "unknown-user", "Pesan Username atau password salah", "XFail DEF-02"],
["LGN-006", "Password salah", "testuser / WrongPassword!", "Pesan Username atau password salah", "XFail DEF-01"],
["LGN-007", "SQL injection", "' OR 1=1 --", "Tetap tidak login", "Pass"],
["LGN-008", "Masking password", "Inspeksi atribut", "type=password", "Pass"],
]
register_rows = [
["REG-001", "Elemen halaman tampil", "Buka /register.php", "Seluruh field dan tombol terlihat", "Pass"],
["REG-002", "Navigasi ke login", "Klik Login", "URL login.php dan judul Sign-In", "Pass"],
["REG-003", "Field kosong", "Submit semua kosong", "Pesan Data tidak boleh kosong", "Pass"],
["REG-004", "Password berbeda", "Password123! / Different123!", "Pesan Password tidak sama", "Pass"],
["REG-005", "Registrasi valid", "Budi Santoso", "Redirect dan name/username/email tersimpan benar", "XFail DEF-04"],
["REG-006", "Username duplikat", "username testuser", "Ditolak, jumlah record tetap 1", "XFail DEF-03/05"],
["REG-007", "Email invalid", "bukan-email", "HTML5 browser menahan submit", "Pass"],
["REG-008", "Masking password", "Inspeksi dua field", "Keduanya type=password", "Pass"],
["REG-009", "Input hanya spasi", "Semua field berisi spasi", "Pesan data kosong", "Pass"],
]

doc.add_heading("5.1 Testcase Login", level=2)
add_table(doc, ["ID", "Skenario", "Langkah/Data", "Expected Result", "Status"], login_rows, font_size=7.7)
doc.add_heading("5.2 Testcase Register", level=2)
add_table(doc, ["ID", "Skenario", "Langkah/Data", "Expected Result", "Status"], register_rows, font_size=7.7)

doc.add_heading("5.3 Format detail testcase yang digunakan", level=2)
add_table(doc, ["Atribut", "Isi"], [
    ["ID", "Kode unik seperti LGN-004 atau REG-006."],
    ["Tujuan", "Perilaku spesifik yang ingin diverifikasi."],
    ["Precondition", "Server aktif, database tersedia, dan data seed telah dibuat."],
    ["Steps", "Buka halaman, isi field, klik tombol, lalu tunggu kondisi tertentu."],
    ["Test data", "Nilai input yang eksplisit dan dapat diulang."],
    ["Expected result", "URL, pesan, elemen, session, atau record database yang harus muncul."],
    ["Evidence", "JUnit, log server, dan screenshot saat gagal."],
], font_size=9)

# 6 Implementation

doc.add_heading("6. Implementasi Selenium", level=1)
doc.add_heading("6.1 Struktur folder final", level=2)
add_code(doc, """quiz-login-register-selenium/
├── .github/workflows/selenium.yml
├── db/quiz_pengupil.sql
├── tests/
│   ├── page_objects/
│   │   ├── base_page.py
│   │   ├── login_page.py
│   │   └── register_page.py
│   ├── stubs/index.php
│   ├── conftest.py
│   ├── test_login.py
│   └── test_register.py
├── scripts/start-local.sh
├── scripts/stop-local.sh
├── login.php
├── register.php
├── koneksi.php
├── requirements.txt
├── pytest.ini
├── TEST_CASES.md
└── README.md""")

doc.add_heading("6.2 Dependency", level=2)
add_code(doc, (ROOT / "requirements.txt").read_text(), "requirements.txt")

doc.add_heading("6.3 Konfigurasi pytest", level=2)
add_code(doc, (ROOT / "pytest.ini").read_text(), "pytest.ini")

doc.add_heading("6.4 Koneksi database yang dapat dipakai lokal dan CI", level=2)
doc.add_paragraph("koneksi.php mempertahankan default XAMPP, tetapi dapat dioverride oleh environment variables pada GitHub Actions. Ini adalah perubahan testability, bukan perubahan aturan bisnis login/register.")
add_code(doc, (ROOT / "koneksi.php").read_text(), "koneksi.php")

doc.add_heading("6.5 Stub index", level=2)
add_code(doc, (ROOT / "tests/stubs/index.php").read_text(), "tests/stubs/index.php")

doc.add_heading("6.6 Fixture dan WebDriver", level=2)
doc.add_paragraph("conftest.py menangani koneksi database, reset data, pencarian ChromeDriver, opsi headless, lifecycle browser, dan screenshot otomatis saat gagal.")

# 7 local

doc.add_heading("7. Menjalankan Pengujian Lokal", level=1)
doc.add_heading("7.1 Import database", level=2)
add_code(doc, "mysql -u root -p < db/quiz_pengupil.sql")
doc.add_paragraph("Pada XAMPP default tanpa password, gunakan mysql -u root < db/quiz_pengupil.sql.")

doc.add_heading("7.2 Set environment", level=2)
add_code(doc, """$env:DB_HOST='127.0.0.1'
$env:DB_PORT='3306'
$env:DB_USER='root'
$env:DB_PASSWORD=''
$env:DB_NAME='quiz_pengupil'
$env:BASE_URL='http://127.0.0.1:8000'""", "PowerShell")

doc.add_heading("7.3 Jalankan server", level=2)
add_code(doc, """Copy-Item tests/stubs/index.php index.php
php -d display_errors=0 -S 127.0.0.1:8000""", "Terminal 1 - Windows")
add_code(doc, "./scripts/start-local.sh", "Terminal 1 - Linux/macOS")

doc.add_heading("7.4 Jalankan Selenium", level=2)
add_code(doc, "pytest -ra --junitxml=artifacts/junit.xml", "Terminal 2")

doc.add_heading("7.5 Hasil yang diharapkan", level=2)
add_status_box(doc, "Expected result", "17 testcase terkoleksi. Selama defect asal belum diperbaiki, ringkasan ideal adalah 13 passed dan 4 xfailed. Pipeline tetap hijau karena xfail adalah hasil yang diharapkan, bukan error infrastruktur.")

doc.add_heading("7.6 Troubleshooting lokal", level=2)
add_table(doc, ["Masalah", "Penyebab umum", "Solusi"], [
    ["Connection refused MySQL", "MySQL belum aktif atau port berbeda", "Aktifkan MySQL/XAMPP dan cek DB_PORT."],
    ["mysqli not found", "Ekstensi PHP belum aktif", "Aktifkan extension=mysqli pada php.ini."],
    ["ChromeDriver version mismatch", "Driver tidak cocok dengan Chrome", "Perbarui driver atau biarkan Selenium Manager menangani driver."],
    ["index.php 404", "Stub belum disalin", "Salin tests/stubs/index.php ke root sebagai index.php."],
    ["Port 8000 dipakai", "Server lain aktif", "Hentikan proses atau ubah BASE_URL dan port php -S."],
    ["Test saling memengaruhi", "Fixture DB tidak berjalan", "Pastikan kredensial DB benar agar reset_database dapat TRUNCATE."],
], font_size=8.3)

# 8 CI

doc.add_heading("8. CI/CD GitHub Actions", level=1)
doc.add_heading("8.1 Trigger", level=2)
add_bullets(doc, [
    "push ke branch main/master",
    "pull request ke main/master",
    "workflow_dispatch untuk menjalankan manual dari tab Actions",
])

doc.add_heading("8.2 Urutan pipeline", level=2)
add_numbered(doc, [
    "Membuat runner Ubuntu 24.04.",
    "Membuat MySQL 8.4 service container dan menunggu health check.",
    "Checkout source code.",
    "Menyiapkan Python 3.12 dan memasang requirements.",
    "Menampilkan versi Chrome, ChromeDriver, dan PHP sebagai bukti environment.",
    "Mengimpor skema database.",
    "Menyalin stub index.php.",
    "Menjalankan PHP built-in server di 127.0.0.1:8000.",
    "Menjalankan pytest dan membuat JUnit XML.",
    "Mengunggah artifact walaupun test gagal, sehingga log dan screenshot tetap tersedia.",
])

doc.add_heading("8.3 Mengapa service container", level=2)
doc.add_paragraph("MySQL service container memberi database baru untuk setiap job. Karena job berjalan langsung pada runner Ubuntu, port service dipetakan ke 127.0.0.1:3306. Setelah job selesai, GitHub menghapus container secara otomatis.")

doc.add_heading("8.4 Cara membaca hasil Actions", level=2)
add_numbered(doc, [
    "Buka repository hasil di GitHub.",
    "Pilih tab Actions.",
    "Buka workflow Selenium Login Register Tests.",
    "Pastikan seluruh step hijau.",
    "Unduh artifact selenium-test-evidence untuk melihat junit.xml, php-server.log, dan screenshot kegagalan.",
])

# 9 upload

doc.add_heading("9. Mengunggah Repository", level=1)
doc.add_heading("9.1 Buat repository kosong", level=2)
add_numbered(doc, [
    "Masuk ke GitHub dan pilih New repository.",
    "Nama yang disarankan: quiz-login-register-selenium.",
    "Pilih Public agar dosen dapat membukanya.",
    "Jangan membuat README baru karena paket sudah memiliki README.",
])

doc.add_heading("9.2 Push paket", level=2)
add_code(doc, """git init
git add .
git commit -m "Add Selenium tests and GitHub Actions pipeline"
git branch -M main
git remote add origin https://github.com/USERNAME_ANDA/quiz-login-register-selenium.git
git push -u origin main""")

doc.add_heading("9.3 Link yang dicantumkan", level=2)
add_status_box(doc, "Repository hasil", "https://github.com/USERNAME_ANDA/quiz-login-register-selenium", "D9EAF7")
doc.add_paragraph("Saya tidak dapat membuat repository langsung pada akun GitHub pengguna tanpa akses akun. Karena itu paket ini sudah berupa Git repository lokal dengan commit awal; Anda hanya perlu membuat repo kosong dan menjalankan perintah push di atas.")

# 10 defects

doc.add_heading("10. Temuan Defect", level=1)
defects = [
("DEF-01", "High", "Password salah tidak menampilkan pesan", "Login dengan testuser dan WrongPassword!", "Tetap di halaman tanpa alert", "Tampilkan pesan generik agar tidak membocorkan apakah username ada."),
("DEF-02", "Medium", "Pesan username tidak ditemukan salah konteks", "Login dengan username yang tidak ada", "Register User Gagal !!", "Ganti dengan Username atau password salah !!."),
("DEF-03", "Critical", "Pengecekan duplikasi memakai name", "Register username testuser dengan name berbeda", "Duplikat dapat tersimpan", "Panggil cek_nama($username, $con) dan ubah nama fungsi menjadi cek_username."),
("DEF-04", "Critical", "Variabel $nama undefined pada INSERT", "Register data valid", "Name kosong atau query gagal", "Gunakan $name dan prepared statement."),
("DEF-05", "High", "Tidak ada unique constraint", "Insert username/email sama dua kali", "Database menerima duplikat", "Tambahkan UNIQUE(username) dan bila diperlukan UNIQUE(email)."),
("DEF-06", "Medium", "Tujuan redirect tidak ada", "Login/register sukses", "index.php 404", "Tambahkan dashboard atau ubah redirect; selama test digunakan stub."),
]
add_table(doc, ["ID", "Severity", "Judul", "Reproduksi", "Actual", "Rekomendasi"], [list(x) for x in defects], font_size=7.2)

# 11 validation

doc.add_heading("11. Validasi dan Checklist Pengumpulan", level=1)
doc.add_heading("11.1 Validasi teknis yang dilakukan pada paket", level=2)
add_table(doc, ["Pemeriksaan", "Status", "Keterangan"], [
    ["PHP lint", "PASS", "login.php, register.php, koneksi.php, dan stub tidak memiliki syntax error."],
    ["Python compile", "PASS", "Seluruh file tests dapat dikompilasi."],
    ["Pytest collection", "PASS", "17 testcase berhasil ditemukan."],
    ["Workflow YAML parse", "PASS", "Struktur job dan steps dapat dibaca."],
    ["Git repository", "PASS", "Repository lokal sudah di-init dan memiliki commit."],
    ["Full E2E lokal", "Belum dijalankan di container pembuatan", "Environment pembuatan tidak menyediakan MySQL server dan ChromeDriver. Full run dirancang untuk GitHub Actions yang menyediakan keduanya."],
], font_size=8.4)

doc.add_heading("11.2 Checklist sebelum submit", level=2)
check_items = [
    "Isi Nama, NIM, dan Kelas pada halaman sampul.",
    "Upload ZIP ke repository GitHub public milik Anda.",
    "Ganti USERNAME_ANDA pada README dan PDF.",
    "Pastikan tab Actions menunjukkan workflow hijau.",
    "Unduh artifact dan simpan screenshot hasil Actions bila dosen meminta bukti tambahan.",
    "Kumpulkan PDF final, bukan DOCX.",
    "Pastikan link repository dapat dibuka tanpa login.",
]
for item in check_items:
    doc.add_paragraph("☐ " + item)

# 12 conclusion

doc.add_heading("12. Kesimpulan", level=1)
doc.add_paragraph("Solusi ini memenuhi unsur utama tugas: penyusunan testcase login dan register, script Selenium, penggunaan Driver dan Stub, integrasi CI/CD GitHub Actions, serta dokumentasi PDF. Suite tidak hanya memeriksa positive flow, tetapi juga validasi kosong, mismatch password, duplikasi, format email, masking password, navigasi, session, database persistence, dan percobaan SQL injection. Known defects ditulis sebagai expected failure agar laporan tetap jujur dan pipeline tetap dapat dipakai sebagai regression baseline.")

# References

doc.add_heading("Referensi", level=1)
refs = [
"Repository objek uji: https://github.com/hermanka/quiz-pengupil",
"Selenium Manager: https://www.selenium.dev/documentation/selenium_manager/",
"Selenium Chrome documentation: https://www.selenium.dev/documentation/webdriver/browsers/chrome/",
"GitHub Actions service containers: https://docs.github.com/actions/use-cases-and-examples/using-containerized-services/about-service-containers",
"GitHub runner image inventory: https://github.com/actions/runner-images/blob/main/images/ubuntu/Ubuntu2404-Readme.md",
]
for x in refs:
    doc.add_paragraph(x)

# Appendices

h = doc.add_heading("Lampiran A - GitHub Actions Workflow", level=1)
h.paragraph_format.page_break_before = True
add_code(doc, compact_code((ROOT / ".github/workflows/selenium.yml").read_text()), ".github/workflows/selenium.yml")

h = doc.add_heading("Lampiran B - Fixture dan WebDriver", level=1)
h.paragraph_format.page_break_before = True
add_code(doc, compact_code((ROOT / "tests/conftest.py").read_text()), "tests/conftest.py")
doc.add_heading("Catatan Lampiran B", level=2)
for item in [
    "reset_database dijalankan otomatis sebelum setiap testcase agar test independen dan dapat diulang.",
    "Driver dicari dari CHROMEDRIVER_PATH, PATH, atau CHROMEWEBDRIVER; pada runner GitHub, ChromeDriver sudah tersedia.",
    "Screenshot disimpan hanya ketika testcase gagal sehingga artifact CI tetap ringkas tetapi bukti kegagalan tersedia.",
]:
    doc.add_paragraph(item, style="List Bullet")

h = doc.add_heading("Lampiran C - Script Test Login", level=1)
h.paragraph_format.page_break_before = True
add_code(doc, compact_code((ROOT / "tests/test_login.py").read_text()), "tests/test_login.py")

h = doc.add_heading("Lampiran D - Script Test Register", level=1)
h.paragraph_format.page_break_before = True
add_code(doc, compact_code((ROOT / "tests/test_register.py").read_text()), "tests/test_register.py")

h = doc.add_heading("Lampiran E - Page Object", level=1)
h.paragraph_format.page_break_before = True
for f in ["base_page.py", "login_page.py", "register_page.py"]:
    add_code(doc, compact_code((ROOT / "tests/page_objects" / f).read_text()), f"tests/page_objects/{f}")

doc.add_heading("Lampiran F - Perintah Cepat", level=2)
add_code(doc, r"""# Instalasi
python -m venv .venv
# Windows: .venv\Scripts\activate
# Linux/macOS: source .venv/bin/activate
pip install -r requirements.txt

# Menjalankan aplikasi dan test
php -S 127.0.0.1:8000
pytest -v --junitxml=artifacts/junit.xml

# Upload repository
git init
git add .
git commit -m \"Add Selenium tests and GitHub Actions pipeline\"
git branch -M main
git remote add origin https://github.com/USERNAME_ANDA/quiz-login-register-selenium.git
git push -u origin main""", "Ringkasan perintah")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
